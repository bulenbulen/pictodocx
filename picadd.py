from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -*- coding: utf-8 -*-
#在第n行写入功能名称，并添加相应的图片n张
#doc文档路径，图片目录路径，图片名称，指定行数（默认为0代表不指定行直接add到最后），modpathnamexx-xx-xx-xx
def picaddname(docfile,picfile,picname,rownum=0,modpathname=None):
    #docfile = r'D:/DATA/doctest/test.docx'
    #picfile = r'D:/DATA/doctest'
    modpathname = modpathname if modpathname else picname
    #空文档就加第一段落，后续段落就直接修改内容
    doc= Document(docfile)
    if len(doc.paragraphs) == 0:
        doc.add_paragraph()
    doc.paragraphs[0].text = '说明'  
        #第一段标题居中
    doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    

    
    #0代表未指定行直接add到最后，1代表第一段落，2代表第二段落，以此类推
    if rownum == 0:
        #后续的文字居左，图片居中
        title_para=doc.add_paragraph(modpathname)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        #图片
        picf_para = doc.add_paragraph()
        picf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picf_para.add_run().add_picture(rf'{picfile}/{picname}.png')

        current_para = picf_para
        for i in range(2, 10):
            img_path = rf'{picfile}/{picname}{i}.png'
            try:
                new_para = doc.add_paragraph()
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_para.add_run().add_picture(img_path)
                current_para._element.addnext(new_para._element)
                current_para = new_para
            except Exception:
                break
    #不为0代表指定段落
    elif rownum > 0:
        # 确保目标段落存在；如果不存在，就补足空段落
        while len(doc.paragraphs) <= rownum:
            doc.add_paragraph()
 # 目标段落
        target_para = doc.paragraphs[rownum] 

        # 在第 n 段后插入标题段落，而不是直接改写该段落
        title_para = doc.add_paragraph(modpathname)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 将标题段落插到目标段落后面
        target_para._element.addnext(title_para._element)

        # 先插入第一张图片
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.add_run().add_picture(rf'{picfile}/{picname}.png')
        title_para._element.addnext(img_para._element)

        current_para = img_para

        # 继续插入后续图片
        for i in range(2, 10):
            img_path = rf'{picfile}/{picname}{i}.png'
            try:
                new_para = doc.add_paragraph()
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_para.add_run().add_picture(img_path)
                current_para._element.addnext(new_para._element)
                current_para = new_para
            except Exception:
                break

    doc.save(docfile)


    
'''#print('ok')
    #try:
        #doc.paragraphs[12].text='000'
        #doc.save(newdpath)
    #except:
        #print('ok')
        
#filepath = r'D:/DATA/doctest/test.docx'
#picfiles = r'D:/DATA/doctest'
#picaddname(filepath, picfiles, '业务员资金占用分析表')
#picaddname(filepath, picfiles, '总账自定义报表',1)'''

 